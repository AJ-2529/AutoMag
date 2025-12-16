# backend/fill_template.py

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement, ns
from backend.headings import normalize_heading
import re


def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(ns.qn("w:val"), "single")
        element.set(ns.qn("w:sz"), "8")
        element.set(ns.qn("w:space"), "0")
        element.set(ns.qn("w:color"), "000000")
        borders.append(element)

    tblPr.append(borders)


def normalize_header_cell(text):
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r"[^\w\s]", "", text)
    text = re.sub(r"\s+", " ", text)
    text = text.replace("sl no", "slno")
    return text.strip()


def table_header_key(table):
    return tuple(normalize_header_cell(c) for c in table[0])


def merge_tables(tables):
    merged = {}

    for table in tables:
        header_key = table_header_key(table)
        header = table[0]
        rows = table[1:]

        if header_key not in merged:
            merged[header_key] = {
                "header": header,
                "rows": []
            }

        merged[header_key]["rows"].extend(rows)

    return merged.values()


def insert_after(paragraph, element):
    paragraph._p.addnext(
        element._p if hasattr(element, "_p") else element._tbl
    )


def fill_template(template_path, output_path, doc1, doc2):
    doc = Document(template_path)

    for para in doc.paragraphs:
        key = normalize_heading(para.text)
        if not key:
            continue

        cursor = para

        # ======================================================
        # TECHNICAL ARTICLES (STABLE VERSION)
        # ======================================================
        if key == "TECHNICAL ARTICLES:":
            src = doc1  # AI content only in doc1

            if src[key]["text"]:
                # -------- HEADING --------
                heading_para = doc.add_paragraph()
                run = heading_para.add_run(src[key]["text"][0])
                run.bold = True
                run.font.size = Pt(14)
                insert_after(cursor, heading_para)
                cursor = heading_para

                # -------- IMAGE (REDUCED SIZE) --------
                if src[key]["images"]:
                    img_para = doc.add_paragraph()
                    img_para.add_run().add_picture(
                        src[key]["images"][0],
                        width=Inches(3.0)  # 🔹 reduced size (safe)
                    )
                    insert_after(cursor, img_para)
                    cursor = img_para

                # -------- PARAGRAPHS --------
                for line in src[key]["text"][1:]:
                    p = doc.add_paragraph(line)
                    insert_after(cursor, p)
                    cursor = p

            continue

        # ======================================================
        # NORMAL TEXT SECTIONS
        # ======================================================
        for src in (doc1, doc2):
            for line in src[key]["text"]:
                p = doc.add_paragraph(line)
                insert_after(cursor, p)
                cursor = p

        # ======================================================
        # MERGED TABLES (UNCHANGED & STABLE)
        # ======================================================
        all_tables = doc1[key]["tables"] + doc2[key]["tables"]
        merged_tables = merge_tables(all_tables)

        for t in merged_tables:
            header = t["header"]
            rows = t["rows"]

            table = doc.add_table(
                rows=1 + len(rows),
                cols=len(header)
            )

            # Header
            for c, cell in enumerate(header):
                table.rows[0].cells[c].text = str(cell)

            # Rows
            for r, row in enumerate(rows, start=1):
                for c, cell in enumerate(row):
                    table.rows[r].cells[c].text = str(cell or "")

            add_table_borders(table)
            cursor._p.addnext(table._tbl)

    doc.save(output_path)
